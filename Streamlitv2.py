import os
import streamlit as st
from langchain.prompts import ChatPromptTemplate
from langchain.schema.output_parser import StrOutputParser
from langchain_openai import ChatOpenAI
from docx import Document
from io import BytesIO
from pdfminer.high_level import extract_text as extract_text_from_pdf
from docx import Document as DocxReader

# =========================
# CONFIGURATION
# =========================
st.set_page_config(page_title="AI Cover Letter Generator", page_icon="✉️", layout="centered")
st.title("✉️ AI Cover Letter Generator (ChatGPT API)")

# OpenAI API key
# Set API key from Streamlit secrets
if "OPENAI_API_KEY" in st.secrets:
    os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]
else:
    st.error("❌ OPENAI_API_KEY not found in Streamlit secrets. Please add it in Settings → Secrets.")
    st.stop()

# =========================
# FUNCTIONS
# =========================
def extract_text(file):
    """Extract text from PDF, DOCX, or plain text"""
    if file.type == "application/pdf":
        return extract_text_from_pdf(file)
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = DocxReader(file)
        return "\n".join([p.text for p in doc.paragraphs])
    elif file.type.startswith("text/"):
        return file.read().decode("utf-8")
    else:
        st.error("Unsupported file type.")
        st.stop()

def generate_cover_letter(topic, resume_text, job_text, tone):
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.7)
    prompt = ChatPromptTemplate.from_template(
        """
        Based on the resume and job description below, write a tailored cover letter for the role of {topic}.
        Tone: {tone}
        Keep it professional, under 50 lines, and aligned with the job requirements.

        === RESUME ===
        {resume}

        === JOB DESCRIPTION ===
        {job_desc}
        """
    )
    output_parser = StrOutputParser()
    chain = prompt | llm | output_parser
    return chain.invoke({
        "topic": topic,
        "resume": resume_text,
        "job_desc": job_text,
        "tone": tone
    })

def save_to_docx(content, topic):
    doc = Document()
    doc.add_heading(f"Cover Letter – {topic}", level=1)
    doc.add_paragraph(content)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# =========================
# STREAMLIT UI
# =========================
st.subheader("Upload your Resume and Enter Job Details")

resume_file = st.file_uploader("📄 Upload Resume (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
job_title = st.text_input("💼 Job Title", placeholder="e.g. Mechanical Design Engineer")
job_description = st.text_area("📝 Paste Job Description", placeholder="Paste job description here...", height=200)

tone = st.selectbox("🎨 Select Cover Letter Tone", ["Formal", "Enthusiastic", "Confident", "Friendly"])

if st.button("🚀 Generate Cover Letter"):
    if not resume_file or not job_title or not job_description.strip():
        st.warning("Please upload resume, enter job title, and paste job description.")
        st.stop()

    with st.spinner("💬 Generating your cover letter..."):
        resume_text = extract_text(resume_file)
        response = generate_cover_letter(job_title, resume_text, job_description, tone)
        docx_file = save_to_docx(response, job_title)

    st.success("✅ Cover letter generated!")
    st.write("### 📜 Preview:")
    st.markdown(response)

    st.download_button(
        label="⬇️ Download Cover Letter (.docx)",
        data=docx_file,
        file_name=f"cover_letter_{job_title.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

st.caption("Developed with ❤️ using Streamlit + ChatGPT API")
