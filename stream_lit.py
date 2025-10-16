import os
import streamlit as st
from pdf2image import convert_from_path
import pytesseract
from langchain.prompts import ChatPromptTemplate
from langchain.schema.output_parser import StrOutputParser
from langchain_openai import ChatOpenAI
from docx import Document
from tempfile import NamedTemporaryFile

# =========================
# CONFIGURATION
# =========================
st.set_page_config(page_title="AI Cover Letter Generator", page_icon="✉️", layout="centered")
st.title("✉️ AI Cover Letter Generator (ChatGPT API)")

# Set your OpenAI API key
os.environ["OPENAI_API_KEY"] = "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"

# Tesseract + Poppler setup (update paths if needed)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH = r"D:\CV writer ai\Release-25.07.0-0\poppler-25.07.0\Library\bin"

# =========================
# FUNCTIONS
# =========================
def check_poppler():
    if not os.path.exists(os.path.join(POPPLER_PATH, "pdftoppm.exe")):
        st.error(f"❌ Poppler not found at `{POPPLER_PATH}`. Please install it.")
        st.stop()

def ocr_pdf(uploaded_file):
    """Extract text from a PDF using OCR"""
    with NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(uploaded_file.read())
        tmp_path = tmp.name

    check_poppler()
    images = convert_from_path(tmp_path, poppler_path=POPPLER_PATH)
    text = ""
    for image in images:
        text += pytesseract.image_to_string(image) + "\n"
    return text.replace("\n", " ").strip()

def generate_cover_letter(topic, resume_text, job_text):
    """Generate cover letter using ChatGPT"""
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.7)
    prompt = ChatPromptTemplate.from_template(
        """
        Based on the resume and job description below, write a tailored cover letter for the role of {topic}.
        Keep it professional, under 50 lines, and aligned with the job requirements.

        === RESUME ===
        {resume}

        === JOB DESCRIPTION ===
        {job_desc}
        """
    )
    output_parser = StrOutputParser()
    chain = prompt | llm | output_parser
    return chain.invoke({
        "topic": topic,
        "resume": resume_text,
        "job_desc": job_text
    })

def save_to_docx(content, topic):
    """Save the generated cover letter to a DOCX file"""
    doc = Document()
    doc.add_heading(f"Cover Letter – {topic}", level=1)
    doc.add_paragraph(content)
    output_path = f"generated_cover_letter_{topic.replace(' ', '_')}.docx"
    doc.save(output_path)
    return output_path

# =========================
# STREAMLIT UI
# =========================
st.subheader("Upload your Resume and Enter Job Details")

resume_file = st.file_uploader("📄 Upload your Resume (PDF)", type=["pdf"])
job_title = st.text_input("💼 Job Title", placeholder="e.g. Mechanical Design Engineer")

job_description = st.text_area(
    "📝 Paste Job Description Text",
    placeholder="Paste the job description here...",
    height=200
)

if st.button("🚀 Generate Cover Letter"):
    if not resume_file or not job_title or not job_description.strip():
        st.warning("Please upload your resume, enter the job title, and paste the job description.")
        st.stop()

    with st.spinner("🔍 Extracting text from resume..."):
        resume_text = ocr_pdf(resume_file)

    with st.spinner("💬 Generating your tailored cover letter..."):
        response = generate_cover_letter(job_title, resume_text, job_description)

    with st.spinner("📁 Saving to DOCX..."):
        output_path = save_to_docx(response, job_title)

    st.success("✅ Cover letter generated successfully!")
    st.write("### 📜 Preview:")
    st.markdown(response)

    with open(output_path, "rb") as f:
        st.download_button(
            label="⬇️ Download Cover Letter (.docx)",
            data=f,
            file_name=os.path.basename(output_path),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

st.caption("Developed with ❤️ using Streamlit + ChatGPT API")
