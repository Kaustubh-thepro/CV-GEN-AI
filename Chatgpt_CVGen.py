import os
from pdf2image import convert_from_path
import pytesseract
from langchain.prompts import ChatPromptTemplate
from langchain.schema.output_parser import StrOutputParser
from langchain_openai import ChatOpenAI
from rich.console import Console

console = Console()
os.environ["OPENAI_API_KEY"] = "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"  # Replace with yours

# =======================
# Step 1: User Input
# =======================
console.print("[bold cyan]--- Cover Letter Generator (ChatGPT API) ---[/bold cyan]\n")

resume_path = input("Enter path to your Resume PDF: ").strip()
jd_path = input("Enter path to the Job Description PDF: ").strip()
topic = input("Enter the Job Title (e.g. Mechanical Design Engineer): ").strip()

# =======================
# Step 2: OCR Setup
# =======================
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH = r"D:\CV writer ai\Release-25.07.0-0\poppler-25.07.0\Library\bin"

def check_poppler():
    """Ensure Poppler is installed correctly."""
    if not os.path.exists(os.path.join(POPPLER_PATH, "pdftoppm.exe")):
        console.print(f"[red]ERROR:[/red] Poppler not found at {POPPLER_PATH}")
        console.print("Please install Poppler for Windows and make sure 'bin' path is correct.")
        console.print("Download: https://github.com/oschwartz10612/poppler-windows/releases/")
        exit(1)

def ocr_pdf(path):
    """Extract text from a PDF using OCR."""
    console.print(f"[yellow]Extracting text from {os.path.basename(path)}...[/yellow]")
    check_poppler()
    images = convert_from_path(path, poppler_path=POPPLER_PATH)
    text = ""
    for i, image in enumerate(images):
        page_text = pytesseract.image_to_string(image)
        text += page_text + "\n"
    return text

def clean_text(text):
    return text.replace('\n', ' ').strip()

# =======================
# Step 3: Extract Text
# =======================
resume_text = clean_text(ocr_pdf(resume_path))
job_text = clean_text(ocr_pdf(jd_path))

console.print("[green]Text extraction completed successfully![/green]\n")

# =======================
# Step 4: Setup ChatGPT (OpenAI)
# =======================
if not os.getenv("OPENAI_API_KEY"):
    console.print("[red]ERROR:[/red] OPENAI_API_KEY not set. Please export it before running.")
    exit(1)

llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.7)

prompt = ChatPromptTemplate.from_template(
    """
    Based on the resume and job description below, write a tailored cover letter for the role of {topic}.
    Keep it under 50 lines, professional, and aligned with the job requirements.

    === RESUME ===
    {resume}

    === JOB DESCRIPTION ===
    {job_desc}
    """
)

output_parser = StrOutputParser()
chain = prompt | llm | output_parser

# =======================
# Step 5: Generate Cover Letter
# =======================
console.print("[cyan]Generating your custom cover letter using ChatGPT...[/cyan]")
response = chain.invoke({
    "topic": topic,
    "resume": resume_text,
    "job_desc": job_text
})

# =======================
# Step 6: Output Result (DOCX Export)
# =======================
from docx import Document

output_file = "generated_cover_letter_chatgpt.docx"

# Create a new Word document
doc = Document()
doc.add_heading("Cover Letter", level=1)
doc.add_paragraph(response)

# Save to DOCX file
doc.save(output_file)

console.print(f"\n[bold green]✅ Cover letter generated successfully![/bold green]")
console.print(f"Saved to: [underline]{output_file}[/underline]\n")
console.print("[bold white]Preview:[/bold white]\n")
console.print(response)

